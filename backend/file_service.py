"""CV parsing (PDF/DOCX/TXT) and document generation (PDF/DOCX) — v3 professional template.

Matches the reference template:
- Centered name + centered contact line
- Dark-grey ALL-CAPS section headings with thin horizontal rule beneath
- 4-column categorized Core Skills grid
- Job: title (left) + dates (right) on one row; bold company on next line; location below
- Education: degree (left) + year (right) on one row; institution | location on next line
"""
import io
from typing import Any, List
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor, black
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY


HEADING_COLOR = HexColor("#1F2937")   # deep charcoal
RULE_COLOR = HexColor("#9CA3AF")      # mid grey for the thin rule
BODY_COLOR = HexColor("#0A0A0A")

DOCX_HEADING_RGB = RGBColor(0x1F, 0x29, 0x37)
DOCX_BODY_RGB = RGBColor(0x0A, 0x0A, 0x0A)


# ===================== Parsing =====================

def parse_cv_bytes(file_bytes: bytes, filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    if name.endswith(".docx"):
        return _parse_docx(file_bytes)
    if name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    raise ValueError(f"Unsupported file type: {filename}. Use PDF, DOCX, or TXT.")


def _parse_pdf(b: bytes) -> str:
    reader = PdfReader(io.BytesIO(b))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
    return "\n".join(parts).strip()


def _parse_docx(b: bytes) -> str:
    doc = Document(io.BytesIO(b))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


# ===================== Normalizers =====================

def _norm_cert(c: Any) -> dict:
    if isinstance(c, dict):
        return {
            "name": (c.get("name") or "").strip(),
            "institute": (c.get("institute") or c.get("issuer") or "").strip(),
            "year": (c.get("year") or c.get("date") or "").strip(),
        }
    return {"name": str(c).strip(), "institute": "", "year": ""}


def _norm_pub(p: Any) -> dict:
    if isinstance(p, dict):
        return {
            "title": (p.get("title") or p.get("name") or "").strip(),
            "publisher": (p.get("publisher") or p.get("journal") or "").strip(),
            "year": (p.get("year") or "").strip(),
            "description": (p.get("description") or "").strip(),
        }
    return {"title": str(p).strip(), "publisher": "", "year": "", "description": ""}


def _skill_groups(cv: dict) -> List[dict]:
    """Return list of {category, items}. Falls back to grouping core_skills under 'Core Competencies'."""
    groups = cv.get("skill_groups") or []
    norm = []
    for g in groups:
        if isinstance(g, dict):
            cat = (g.get("category") or "").strip()
            items = [str(x).strip() for x in (g.get("items") or []) if str(x).strip()]
            if cat and items:
                norm.append({"category": cat, "items": items})
    if norm:
        return norm
    # Backwards-compat: flat core_skills list
    flat = cv.get("core_skills") or []
    flat = [str(x).strip() for x in flat if str(x).strip()]
    if flat:
        return [{"category": "Core Competencies", "items": flat}]
    return []


# ===================== DOCX helpers =====================

def _set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            b = OxmlElement(f"w:{edge}")
            for k, v in kwargs[edge].items():
                b.set(qn(f"w:{k}"), str(v))
            tc_borders.append(b)
    tc_pr.append(tc_borders)


def _section_heading_docx(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = DOCX_HEADING_RGB
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    # Apply a bottom border to act as the rule
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F2937")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _title_with_right_tab(doc: Document, left_text: str, left_bold: bool, right_text: str,
                          font_size: int = 11):
    """Paragraph with a right-aligned tab stop for placing dates/years on the right edge."""
    p = doc.add_paragraph()
    # right tab stop at ~6.5"
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_TAB_ALIGNMENT.RIGHT)
    left_run = p.add_run(left_text)
    left_run.bold = left_bold
    left_run.font.size = Pt(font_size)
    left_run.font.color.rgb = DOCX_BODY_RGB
    if right_text:
        tab_run = p.add_run("\t")
        right_run = p.add_run(right_text)
        right_run.font.size = Pt(font_size)
        right_run.font.color.rgb = DOCX_BODY_RGB
    p.paragraph_format.space_after = Pt(0)
    return p


# ===================== DOCX generation =====================

def generate_cv_docx(cv: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # ---- Header: centered name + contact ----
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv.get("full_name", "").upper())
    name_run.bold = True
    name_run.font.size = Pt(28)
    name_run.font.color.rgb = DOCX_BODY_RGB
    name_p.paragraph_format.space_after = Pt(2)

    contact = cv.get("contact", {}) or {}
    primary = [contact.get(k, "") for k in ("email", "phone", "location", "linkedin")]
    primary = [p for p in primary if p]
    if primary:
        c_p = doc.add_paragraph()
        c_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c_run = c_p.add_run(" | ".join(primary))
        c_run.font.size = Pt(10)
        c_run.font.color.rgb = DOCX_BODY_RGB
        c_p.paragraph_format.space_after = Pt(0)
    if contact.get("website"):
        w_p = doc.add_paragraph()
        w_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        w_run = w_p.add_run(contact["website"])
        w_run.font.size = Pt(10)
        w_run.font.color.rgb = DOCX_BODY_RGB
        w_p.paragraph_format.space_after = Pt(4)

    # ---- Objective ----
    objective = (cv.get("objective") or "").strip()
    if not objective and cv.get("headline"):
        objective = f"To contribute as {cv['headline']}."
    if objective:
        _section_heading_docx(doc, "Objective")
        p = doc.add_paragraph(objective)
        for r in p.runs:
            r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)

    # ---- Professional Summary ----
    if cv.get("professional_summary"):
        _section_heading_docx(doc, "Professional Summary")
        p = doc.add_paragraph(cv["professional_summary"])
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs:
            r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)

    # ---- Core Skills (4-column grid) ----
    groups = _skill_groups(cv)
    if groups:
        _section_heading_docx(doc, "Core Skills")
        ncols = min(4, max(1, len(groups)))
        # spacer paragraph
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sp.paragraph_format.space_after = Pt(0)
        table = doc.add_table(rows=2, cols=ncols)
        table.autofit = False
        # Distribute width
        page_width = Inches(7.0)
        col_w = Inches(7.0 / ncols)
        for col in table.columns:
            col.width = col_w
        # Row 1: bold headers
        for i in range(ncols):
            cell = table.rows[0].cells[i]
            cell.width = col_w
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(groups[i]["category"].upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DOCX_HEADING_RGB
            p.paragraph_format.space_after = Pt(2)
            _set_cell_border(cell,
                             top={"val": "nil"}, left={"val": "nil"},
                             right={"val": "nil"}, bottom={"val": "nil"})
        # Row 2: items
        for i in range(ncols):
            cell = table.rows[1].cells[i]
            cell.width = col_w
            cell.text = ""
            p0 = cell.paragraphs[0]
            for j, item in enumerate(groups[i]["items"]):
                p = p0 if j == 0 else cell.add_paragraph()
                run = p.add_run(item)
                run.font.size = Pt(10)
                run.font.color.rgb = DOCX_BODY_RGB
                p.paragraph_format.space_after = Pt(1)
            _set_cell_border(cell,
                             top={"val": "nil"}, left={"val": "nil"},
                             right={"val": "nil"}, bottom={"val": "nil"})

    # ---- Professional Experience ----
    exp = cv.get("experience") or []
    if exp:
        _section_heading_docx(doc, "Professional Experience")
        for idx, job in enumerate(exp):
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(6 if idx == 0 else 8)
            spacer.paragraph_format.space_after = Pt(0)
            # Row 1: title (left, bold) + dates (right)
            dates = f"{job.get('start_date','')} – {job.get('end_date','')}".strip(" –")
            _title_with_right_tab(doc, job.get("title", ""), True, dates, font_size=11)
            # Row 2: company (bold)
            if job.get("company"):
                cp = doc.add_paragraph()
                cp_run = cp.add_run(job["company"])
                cp_run.bold = True
                cp_run.font.size = Pt(11)
                cp_run.font.color.rgb = DOCX_BODY_RGB
                cp.paragraph_format.space_after = Pt(0)
            # Row 3: location
            if job.get("location"):
                lp = doc.add_paragraph()
                lp_run = lp.add_run(job["location"])
                lp_run.font.size = Pt(11)
                lp_run.font.color.rgb = DOCX_BODY_RGB
                lp.paragraph_format.space_after = Pt(3)
            # Bullets
            for b in job.get("bullets", []):
                bp = doc.add_paragraph(b, style="List Bullet")
                for r in bp.runs:
                    r.font.size = Pt(10.5)
                bp.paragraph_format.space_after = Pt(2)

    # ---- Education ----
    edu = cv.get("education") or []
    if edu:
        _section_heading_docx(doc, "Education")
        for e in edu:
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(4)
            spacer.paragraph_format.space_after = Pt(0)
            # Row 1: degree (left, bold) + year (right)
            year = e.get("end_date", "") or e.get("start_date", "")
            _title_with_right_tab(doc, e.get("degree", ""), True, year, font_size=11)
            # Row 2: institution | location
            inst_parts = [x for x in [e.get("institution", ""), e.get("location", "")] if x]
            if inst_parts:
                ip = doc.add_paragraph()
                ip_run = ip.add_run(" | ".join(inst_parts))
                ip_run.font.size = Pt(11)
                ip_run.font.color.rgb = DOCX_BODY_RGB
                ip.paragraph_format.space_after = Pt(0)
            if e.get("details"):
                dp = doc.add_paragraph()
                dp_run = dp.add_run(e["details"])
                dp_run.font.size = Pt(10.5)
                dp_run.font.color.rgb = DOCX_BODY_RGB
                dp.paragraph_format.space_after = Pt(2)

    # ---- Certifications ----
    certs = [_norm_cert(c) for c in (cv.get("certifications") or [])]
    certs = [c for c in certs if c["name"]]
    if certs:
        _section_heading_docx(doc, "Certifications")
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(4)
        sp.paragraph_format.space_after = Pt(0)
        for c in certs:
            line_text = c["name"]
            if c["institute"]:
                line_text += f" | {c['institute']}"
            _title_with_right_tab(doc, line_text, False, c["year"], font_size=11)

    # ---- Projects ----
    projects = cv.get("projects") or []
    if projects:
        _section_heading_docx(doc, "Projects")
        for proj in projects:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(4)
            sp.paragraph_format.space_after = Pt(0)
            p = doc.add_paragraph()
            t = p.add_run(proj.get("name", ""))
            t.bold = True
            t.font.size = Pt(11)
            t.font.color.rgb = DOCX_BODY_RGB
            p.paragraph_format.space_after = Pt(0)
            if proj.get("description"):
                d = doc.add_paragraph(proj["description"])
                for r in d.runs:
                    r.font.size = Pt(11)
                d.paragraph_format.space_after = Pt(2)

    # ---- Publications ----
    pubs = [_norm_pub(p) for p in (cv.get("publications") or [])]
    pubs = [p for p in pubs if p["title"]]
    if pubs:
        _section_heading_docx(doc, "Publications")
        for pub in pubs:
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(4)
            sp.paragraph_format.space_after = Pt(0)
            meta = ", ".join([x for x in [pub["publisher"], pub["year"]] if x])
            _title_with_right_tab(doc, pub["title"], True, meta, font_size=11)
            if pub["description"]:
                d = doc.add_paragraph(pub["description"])
                for r in d.runs:
                    r.font.size = Pt(11)
                d.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_cover_letter_docx(text: str, candidate_name: str) -> bytes:
    """Render the AI-generated cover letter text as a clean left-aligned business letter."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # The AI returns blocks separated by blank lines. Render each block as a paragraph,
    # preserving internal line-breaks for blocks like the sender header & recipient block.
    blocks = [b.strip("\n") for b in text.replace("\r\n", "\n").split("\n\n") if b.strip()]
    for block in blocks:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lines = block.split("\n")
        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.size = Pt(11)
            run.font.color.rgb = DOCX_BODY_RGB
        p.paragraph_format.space_after = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ===================== PDF generation =====================

def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "name_center": ParagraphStyle("NameC", parent=styles["Title"], fontName="Helvetica-Bold",
                                       fontSize=26, leading=30, textColor=BODY_COLOR,
                                       alignment=TA_CENTER, spaceAfter=2),
        "contact_center": ParagraphStyle("ContactC", parent=styles["Normal"], fontName="Helvetica",
                                          fontSize=10, leading=13, textColor=BODY_COLOR,
                                          alignment=TA_CENTER, spaceAfter=2),
        "section": ParagraphStyle("Section", parent=styles["Heading2"], fontName="Helvetica-Bold",
                                   fontSize=12, leading=15, textColor=HEADING_COLOR, spaceBefore=12,
                                   spaceAfter=0, alignment=TA_LEFT),
        "objective": ParagraphStyle("Objective", parent=styles["Normal"], fontName="Helvetica",
                                     fontSize=11, leading=15, textColor=BODY_COLOR, spaceAfter=2, alignment=TA_LEFT),
        "summary": ParagraphStyle("Summary", parent=styles["Normal"], fontName="Helvetica",
                                   fontSize=11, leading=15, textColor=BODY_COLOR, spaceAfter=2, alignment=TA_JUSTIFY),
        "skill_header": ParagraphStyle("SkillH", parent=styles["Normal"], fontName="Helvetica-Bold",
                                        fontSize=10, leading=13, textColor=HEADING_COLOR, spaceAfter=4, alignment=TA_LEFT),
        "skill_item": ParagraphStyle("SkillI", parent=styles["Normal"], fontName="Helvetica",
                                      fontSize=10, leading=13, textColor=BODY_COLOR, spaceAfter=2, alignment=TA_LEFT),
        "job_title": ParagraphStyle("JT", parent=styles["Normal"], fontName="Helvetica-Bold",
                                     fontSize=11.5, leading=14, textColor=BODY_COLOR, spaceAfter=0, alignment=TA_LEFT),
        "job_date_right": ParagraphStyle("JD", parent=styles["Normal"], fontName="Helvetica",
                                          fontSize=11, leading=14, textColor=BODY_COLOR, spaceAfter=0, alignment=TA_RIGHT),
        "company": ParagraphStyle("Co", parent=styles["Normal"], fontName="Helvetica-Bold",
                                   fontSize=11, leading=14, textColor=BODY_COLOR, spaceAfter=0, alignment=TA_LEFT),
        "location": ParagraphStyle("Loc", parent=styles["Normal"], fontName="Helvetica",
                                    fontSize=11, leading=14, textColor=BODY_COLOR, spaceAfter=4, alignment=TA_LEFT),
        "bullet": ParagraphStyle("Bul", parent=styles["Normal"], fontName="Helvetica",
                                  fontSize=10.5, leading=14, textColor=BODY_COLOR, leftIndent=14,
                                  bulletIndent=2, spaceAfter=2, alignment=TA_JUSTIFY),
        "body": ParagraphStyle("Body", parent=styles["Normal"], fontName="Helvetica",
                                fontSize=11, leading=14, textColor=BODY_COLOR, spaceAfter=3, alignment=TA_LEFT),
    }


def _section(story, s, title):
    story.append(Paragraph(title.upper(), s["section"]))
    story.append(HRFlowable(width="100%", thickness=1.0, color=HEADING_COLOR,
                            spaceBefore=1, spaceAfter=6))


def _two_col_row(left_para, right_para, page_width):
    """Title (left) + Date (right) row using a borderless table."""
    t = Table([[left_para, right_para]], colWidths=[page_width * 0.65, page_width * 0.35])
    t.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
    ]))
    return t


def generate_cv_pdf(cv: dict) -> bytes:
    buf = io.BytesIO()
    LEFT = 0.75 * inch
    RIGHT = 0.75 * inch
    page_w = LETTER[0] - LEFT - RIGHT
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            leftMargin=LEFT, rightMargin=RIGHT,
                            topMargin=0.55 * inch, bottomMargin=0.55 * inch)
    s = _pdf_styles()
    story = []

    # Centered header
    story.append(Paragraph(cv.get("full_name", "").upper(), s["name_center"]))
    contact = cv.get("contact", {}) or {}
    primary = [contact.get(k, "") for k in ("email", "phone", "location", "linkedin")]
    primary = [p for p in primary if p]
    if primary:
        story.append(Paragraph(" | ".join(primary), s["contact_center"]))
    if contact.get("website"):
        story.append(Paragraph(contact["website"], s["contact_center"]))
    story.append(Spacer(1, 6))

    # Objective
    objective = (cv.get("objective") or "").strip()
    if not objective and cv.get("headline"):
        objective = f"To contribute as {cv['headline']}."
    if objective:
        _section(story, s, "Objective")
        story.append(Paragraph(objective, s["objective"]))

    # Summary
    if cv.get("professional_summary"):
        _section(story, s, "Professional Summary")
        story.append(Paragraph(cv["professional_summary"], s["summary"]))

    # Core Skills 4-col grid
    groups = _skill_groups(cv)
    if groups:
        _section(story, s, "Core Skills")
        ncols = min(4, max(1, len(groups)))
        col_w = page_w / ncols
        # build header row and items row
        header_cells = [Paragraph(g["category"].upper(), s["skill_header"]) for g in groups[:ncols]]
        item_cells = []
        for g in groups[:ncols]:
            paras = [Paragraph(it, s["skill_item"]) for it in g["items"]]
            item_cells.append(paras)
        t = Table([header_cells, item_cells], colWidths=[col_w] * ncols)
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
        ]))
        story.append(t)
        story.append(Spacer(1, 2))

    # Experience
    exp = cv.get("experience") or []
    if exp:
        _section(story, s, "Professional Experience")
        for idx, job in enumerate(exp):
            if idx > 0:
                story.append(Spacer(1, 6))
            dates = f"{job.get('start_date','')} – {job.get('end_date','')}".strip(" –")
            left = Paragraph(job.get("title", ""), s["job_title"])
            right = Paragraph(dates, s["job_date_right"])
            story.append(_two_col_row(left, right, page_w))
            if job.get("company"):
                story.append(Paragraph(job["company"], s["company"]))
            if job.get("location"):
                story.append(Paragraph(job["location"], s["location"]))
            else:
                story.append(Spacer(1, 2))
            for b in job.get("bullets", []):
                story.append(Paragraph(b, s["bullet"], bulletText="•"))

    # Education
    edu = cv.get("education") or []
    if edu:
        _section(story, s, "Education")
        for idx, e in enumerate(edu):
            if idx > 0:
                story.append(Spacer(1, 4))
            year = e.get("end_date", "") or e.get("start_date", "")
            left = Paragraph(e.get("degree", ""), s["job_title"])
            right = Paragraph(year, s["job_date_right"])
            story.append(_two_col_row(left, right, page_w))
            inst_parts = [x for x in [e.get("institution", ""), e.get("location", "")] if x]
            if inst_parts:
                story.append(Paragraph(" | ".join(inst_parts), s["location"]))
            if e.get("details"):
                story.append(Paragraph(e["details"], s["body"]))

    # Certifications
    certs = [_norm_cert(c) for c in (cv.get("certifications") or [])]
    certs = [c for c in certs if c["name"]]
    if certs:
        _section(story, s, "Certifications")
        for c in certs:
            left_text = c["name"]
            if c["institute"]:
                left_text += f" | {c['institute']}"
            left = Paragraph(f"<b>{left_text}</b>", s["job_title"])
            right = Paragraph(c["year"], s["job_date_right"])
            story.append(_two_col_row(left, right, page_w))
            story.append(Spacer(1, 2))

    # Projects
    projects = cv.get("projects") or []
    if projects:
        _section(story, s, "Projects")
        for proj in projects:
            story.append(Paragraph(f"<b>{proj.get('name','')}</b>", s["body"]))
            if proj.get("description"):
                story.append(Paragraph(proj["description"], s["body"]))
            story.append(Spacer(1, 2))

    # Publications
    pubs = [_norm_pub(p) for p in (cv.get("publications") or [])]
    pubs = [p for p in pubs if p["title"]]
    if pubs:
        _section(story, s, "Publications")
        for pub in pubs:
            meta = ", ".join([x for x in [pub["publisher"], pub["year"]] if x])
            left = Paragraph(f"<b>{pub['title']}</b>", s["job_title"])
            right = Paragraph(meta, s["job_date_right"])
            story.append(_two_col_row(left, right, page_w))
            if pub["description"]:
                story.append(Paragraph(pub["description"], s["body"]))
            story.append(Spacer(1, 2))

    doc.build(story)
    return buf.getvalue()


def generate_cover_letter_pdf(text: str, candidate_name: str) -> bytes:
    """Render the AI-generated cover letter text as a clean left-aligned business letter."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                            leftMargin=1.0 * inch, rightMargin=1.0 * inch,
                            topMargin=1.0 * inch, bottomMargin=1.0 * inch)
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "CLBody", parent=styles["Normal"], fontName="Helvetica",
        fontSize=11, leading=15, textColor=BODY_COLOR, spaceAfter=10, alignment=TA_LEFT,
    )
    story = []
    blocks = [b.strip("\n") for b in text.replace("\r\n", "\n").split("\n\n") if b.strip()]
    for block in blocks:
        # Preserve intra-block line breaks (sender header & recipient block)
        html = block.replace("\n", "<br/>")
        story.append(Paragraph(html, body_style))
    doc.build(story)
    return buf.getvalue()
